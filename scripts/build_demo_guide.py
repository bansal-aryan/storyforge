from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/aryanbansal/Documents/storyforge/docs/Storyforge_Demo_Video_Guide.docx"
BLUE, TEAL, INK, MUTED, PALE, GOLD, RED = "244B68", "28766E", "17252D", "68767D", "EAF1F4", "C18A35", "A74545"

doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin = section.bottom_margin = Inches(0.72)
section.left_margin = section.right_margin = Inches(0.82)
section.header_distance = section.footer_distance = Inches(0.38)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, before, after in [("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 9, 4)]:
    style = doc.styles[name]
    style.font.name = "Aptos Display"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE if name != "Heading 3" else TEAL)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

def shade(paragraph, fill=PALE, border=TEAL):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), border)
    left.set(qn("w:space"), "8")
    borders.append(left)
    ppr.append(borders)

def add_callout(label, text, fill=PALE, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    shade(p, fill, color)
    r = p.add_run(label.upper() + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.size = Pt(9)
    p.add_run(text)

def add_step(time, title, show, narration, prompt=None, overlay=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(time + "  ")
    r.bold = True; r.font.color.rgb = RGBColor.from_string(GOLD); r.font.size = Pt(11)
    r = p.add_run(title)
    r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE); r.font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("SHOW  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run(show)
    if prompt:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.space_after = Pt(3)
        r = p.add_run("PASTE  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(RED)
        q = p.add_run(prompt); q.italic = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("SAY  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run(narration)
    if overlay:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.space_after = Pt(5)
        r = p.add_run("TEXT  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(GOLD)
        p.add_run(overlay)

header = section.header.paragraphs[0]
header.text = "STORYFORGE  /  DEMO VIDEO PLAYBOOK"
header.runs[0].font.size = Pt(8.5); header.runs[0].font.bold = True
header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("WebMCP Hackathon · Recording guide")
r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(MUTED)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(3)
r = p.add_run("DEMO VIDEO PLAYBOOK")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(GOLD)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Storyforge")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor.from_string(INK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("A precise 2-minute recording plan for showing WebMCP as gameplay—not decoration")
r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(MUTED)

add_callout("Target", "Finish between 1:50 and 2:15. Show the working product in the first 5 seconds. Make one complete WebMCP interaction the centerpiece.", "E8F2EF", TEAL)
doc.add_heading("The single story this video tells", level=1)
doc.add_paragraph("Storyforge gives a human player and AI agents one persistent fantasy world. WebMCP lets the agent read canonical game state, choose a grounded next action, portray an in-world companion, and propose a change for the player to approve. The human remains the author; the agent becomes a real participant.")

doc.add_heading("Prepare this exact recording state", level=1)
for text in [
    "Open https://storyforge-rosy.vercel.app/?demo=1 in a clean browser window at 1920 × 1080 or 1440 × 900.",
    "Start after the intro. Keep the Agent panel open and make sure the WebMCP status says available.",
    "Use a prepared Stage 2–5 save for the party-coordination shot. Recruit at least two companions and leave one ritual incomplete.",
    "Close unrelated tabs, notifications, bookmarks, and developer tools. Set browser zoom to 100%.",
    "Paste commands; do not type them live. Record narration separately if that produces cleaner audio.",
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("Prompts to copy before recording", level=2)
for label, prompt in [
    ("Clipboard 1", "Inspect the battlefield and tell me the next valid move."),
    ("Clipboard 2", "Have Elias scout the West Grove, warn me about the blight, and recommend a safe approach."),
    ("Clipboard 3", "Coordinate the fellowship and propose a plan for this ritual."),
]:
    p = doc.add_paragraph()
    p.add_run(label + ": ").bold = True
    p.add_run(prompt)

doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("THE RUN OF SHOW")
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GOLD)
doc.add_heading("Record in six short clips", level=1)
add_callout("Recording rule", "Finish each clip with one second of stillness. This gives you a clean place for jump cuts and captions.")

add_step("0:00–0:10", "Cold open: the game already works",
         "Begin on the moving game world with the Agent panel open. Move the heir for two seconds; do not show the story intro.",
         "This is Storyforge, a fantasy RPG where humans and AI agents share one persistent, canonical world.",
         overlay="One world · Human decisions · Agent initiative")
add_step("0:10–0:32", "WebMCP reads the live battlefield",
         "Show the agent invoking inspect_battlefield. Hold on Agent Activity, then reveal health, pressure, enemies, recruitment gate, and the shielded boss in the HUD.",
         "The agent does not guess from pixels. Through WebMCP, it reads structured battlefield state: health, realm pressure, enemies, recruitment requirements, objectives, and boss phases.",
         prompt="Inspect the battlefield and tell me the next valid move.",
         overlay="inspect_battlefield → canonical state")
add_step("0:32–0:49", "The agent respects quest rules",
         "Show explain_next_objective in Agent Activity and briefly highlight its recommendation. Do not linger on raw JSON.",
         "It then asks the game for the next valid action. That means it can help without inventing progress or bypassing the quest’s rules.",
         overlay="explain_next_objective → grounded guidance")
add_step("0:49–1:17", "An agent acts as Elias",
         "Trigger propose_companion_action. Show the Elias proposal and the visible Accept and Reject controls. Pause before accepting.",
         "Agents can also inhabit the fellowship. Here, Elias chooses to scout the West Grove and warn me about the blight. The action does not silently rewrite the story—it appears as a proposal for me to review.",
         prompt="Have Elias scout the West Grove, warn me about the blight, and recommend a safe approach.",
         overlay="Agent proposal · Player remains in control")
add_step("1:17–1:30", "The player makes it canon",
         "Click Accept. Show the proposal leave the decision queue and appear in Agent Activity. Let the UI settle for one second.",
         "When I accept, that decision becomes part of the world’s logged history. The same canonical state is now shared by the player, the game, and future agent actions.",
         overlay="Accepted → logged world history")

p = doc.add_paragraph()
r = p.add_run("THE RUN OF SHOW  /  CONTINUED")
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GOLD)
add_step("1:30–1:52", "Show why the party matters",
         "Cut to a prepared later-stage save. Approach a ritual that needs a companion ability. Open Party for one second, then show the locked ritual marker.",
         "Later realms cannot be solved by the heir alone. Rituals require different companion expertise, so party composition and ability readiness materially change what the player can do.",
         overlay="Companion expertise unlocks quest paths")
add_step("1:52–2:12", "Coordinate instead of micromanaging",
         "Invoke propose_battle_plan. Show the plan card assigning different actions. Execute it only if the animation is immediate and clean.",
         "The agent evaluates the current fellowship and proposes a coordinated plan. Companions keep distinct abilities, voices, trust, and autonomy—but consequential choices still come back to the player.",
         prompt="Coordinate the fellowship and propose a plan for this ritual.",
         overlay="propose_battle_plan → multi-companion assignments")
add_step("2:12–2:22", "End on the thesis",
         "Return to a clean wide shot containing the world, party, hotbar, and Agent Activity. Stop before another modal appears.",
         "WebMCP turns the agent from a chatbot beside the game into a grounded participant inside it.",
         overlay="Storyforge · Play with agents, not around them")

doc.add_heading("Full narration, continuous read", level=1)
doc.add_paragraph("This is Storyforge, a fantasy RPG where humans and AI agents share one persistent, canonical world. The agent does not guess from pixels. Through WebMCP, it reads structured battlefield state: health, realm pressure, enemies, recruitment requirements, objectives, and boss phases. It then asks the game for the next valid action, so it can help without inventing progress or bypassing quest rules. Agents can also inhabit the fellowship. Here, Elias chooses to scout the West Grove and warn me about the blight. The action does not silently rewrite the story—it appears as a proposal for me to review. When I accept, that decision becomes part of the world’s logged history. Later realms cannot be solved by the heir alone. Rituals require different companion expertise, so party composition and ability readiness materially change what the player can do. The agent evaluates the current fellowship and proposes a coordinated plan. Companions keep distinct abilities, voices, trust, and autonomy—but consequential choices still come back to the player. WebMCP turns the agent from a chatbot beside the game into a grounded participant inside it.")

doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("RECORDING & EDITING CHECKLIST")
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GOLD)
doc.add_heading("Make the footage judge-friendly", level=1)

groups = [
    ("Before recording", ["The live URL loads without login, warnings, or a black screen.", "WebMCP shows available before the first clip.", "The later-stage save contains at least two recruited companions.", "Narration, game audio, and microphone levels pass a ten-second test.", "Every pasted prompt has been rehearsed and triggers the expected tool."]),
    ("During recording", ["Start on gameplay—not a title card, story crawl, setup screen, or refresh.", "Keep the cursor away from important UI unless it is clicking something.", "Record each section separately with one second of stillness at both ends.", "Redo any clip that stalls; never make judges watch waiting time.", "Show tool names, their visible effect, and player approval in one sequence."]),
    ("In the edit", ["Remove loading, typing, dead air, repeated combat, and failed attempts.", "Use short labels that do not cover Agent Activity or the objective HUD.", "Normalize narration and keep game voices beneath it.", "Use hard cuts or brief dissolves—not decorative transitions.", "Export at 1080p, 30 or 60 fps, H.264, with audible AAC audio."]),
]
for heading, items in groups:
    doc.add_heading(heading, level=2)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run("☐  "); r.font.color.rgb = RGBColor.from_string(TEAL)
        p.add_run(item)

add_callout("Cut first if long", "Remove the separate quest-guidance beat at 0:32–0:49 and mention it during battlefield inspection. Never cut the agent proposal, approval, or coordinated party plan.", "FFF5E5", GOLD)
doc.add_heading("Final submission check", level=1)
for text in [
    "Runtime is under 3:00; target is approximately 2:22.",
    "YouTube visibility matches the official requirement.",
    "The first 15 seconds show the working product and explain the core idea.",
    "The video visibly demonstrates WebMCP tools acting on real game state.",
    "No narration claim exceeds what the deployed build demonstrates.",
    "The submission includes the live URL, repository, testing instructions, and WebMCP details.",
]:
    p = doc.add_paragraph()
    r = p.add_run("☐  "); r.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run(text)
add_callout("Final frame", "Hold the final Storyforge shot for two seconds, then fade to black. Do not add a long credits sequence.", "E8F2EF", TEAL)

doc.core_properties.title = "Storyforge Demo Video Guide"
doc.core_properties.subject = "WebMCP hackathon recording playbook"
doc.core_properties.author = "Storyforge"
doc.save(OUT)
print(OUT)
